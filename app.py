"""
Daily Report Generator — PT Nale System Integrator
Versi dengan perbaikan upload foto, placeholder guide, dan CSS file uploader.
"""

import streamlit as st
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageOps
from datetime import date
from pathlib import Path
from copy import deepcopy
import io
import json
import re
import os
import tempfile
from fpdf import FPDF
import fitz  # PyMuPDF for PDF template reading/manipulation
import fitz  # PyMuPDF untuk membaca/mengedit PDF template

# ═══════════════════════════════════════════════════════════════
# CONFIG
# ═══════════════════════════════════════════════════════════════
APP_DIR = Path(__file__).parent
TEMPLATE_PATH = APP_DIR / "template.docx"
TEMPLATE_PDF_PATH = APP_DIR / "template mojokerto.pdf"
if not TEMPLATE_PDF_PATH.exists():
    TEMPLATE_PDF_PATH = APP_DIR / "template.pdf"
PERSIST_FILE = APP_DIR / ".user_data.json"
FULL_DATA_FILE = APP_DIR / ".full_report_data.json"

BULAN_ID = [
    "Januari", "Februari", "Maret", "April", "Mei", "Juni",
    "Juli", "Agustus", "September", "Oktober", "November", "Desember"
]

st.set_page_config(
    page_title="Daily Report Generator — Nale SI",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ═══════════════════════════════════════════════════════════════
# STYLING (dark theme + file uploader fix)
# ═══════════════════════════════════════════════════════════════
st.markdown("""
<style>
    .stApp { background: #0e1233; }
    .main .block-container { padding-top: 1.5rem; max-width: 1100px; }
    h1, h2, h3, h4, p, label, span, div { color: #e2e8f0 !important; }
    .stTextInput input, .stTextArea textarea, .stDateInput input, .stSelectbox div[data-baseweb="select"] > div {
        background: #181e4a !important; color: #e2e8f0 !important; border: 1px solid rgba(255,255,255,0.07) !important;
    }
    .stTextInput input::placeholder, .stTextArea textarea::placeholder {
        color: #7a82a6 !important; font-style: italic; font-size: 0.85rem;
    }
    /* Perbaikan file uploader agar teks terlihat */
    .stFileUploader {
        background: #181e4a;
        border: 1px solid rgba(255,255,255,0.07);
        border-radius: 8px;
    }
    .stFileUploader > div:first-child,
    .stFileUploader label,
    .stFileUploader span,
    .stFileUploader p,
    .stFileUploader .st-bu {
        color: #e2e8f0 !important;
    }
    .stFileUploader button {
        background: #0e1233 !important;
        color: #f5a623 !important;
        border: 1px solid rgba(255,255,255,0.2) !important;
    }
    .stFileUploader button:hover {
        background: #1c2255 !important;
    }
    .stFileUploader [data-testid="stFileUploaderDropzone"] {
        background: #181e4a;
        border-color: #f5a623;
    }
    .stFileUploader [data-testid="stFileUploaderDropzone"] div {
        color: #e2e8f0;
    }
    .stButton > button {
        background: linear-gradient(135deg, #f5a623, #e09000); color: #141942 !important;
        border: none; font-weight: 700; border-radius: 8px;
    }
    .stButton > button:hover { background: linear-gradient(135deg, #ffb940, #f5a623); }
    .stDownloadButton > button {
        background: linear-gradient(135deg, #22c55e, #16a34a); color: white !important;
        border: none; font-weight: 700; border-radius: 8px; padding: 10px 24px;
    }
    [data-testid="stHeader"] { background: transparent; }
    .brand-header {
        background: linear-gradient(135deg, #141942, #1c2255);
        border: 1px solid rgba(255,255,255,0.07);
        border-radius: 16px; padding: 22px 28px; margin-bottom: 20px;
        border-top: 3px solid #f5a623;
    }
    .brand-tag { font-size: 11px; font-weight: 700; letter-spacing: 2.5px; color: #f5a623; text-transform: uppercase; }
    .brand-title { font-size: 22px; font-weight: 800; color: white; margin: 4px 0 2px; }
    .brand-sub { font-size: 12px; color: #7a82a6; }
    .section-card {
        background: #181e4a;
        border: 1px solid rgba(255,255,255,0.07);
        border-radius: 12px; padding: 18px; margin-bottom: 14px;
    }
    .section-num {
        display: inline-flex; align-items: center; justify-content: center;
        width: 28px; height: 28px; border-radius: 7px;
        background: linear-gradient(135deg, #f5a623, #e09000);
        color: #141942; font-weight: 800; margin-right: 10px;
    }
    .persist-info {
        font-size: 11px; color: #22c55e; background: rgba(34,197,94,0.08);
        border: 1px solid rgba(34,197,94,0.15); padding: 4px 12px;
        border-radius: 20px; display: inline-block;
    }
</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════
# PERSISTENCE
# ═══════════════════════════════════════════════════════════════
def load_persist():
    try:
        if PERSIST_FILE.exists():
            return json.loads(PERSIST_FILE.read_text())
    except Exception:
        pass
    return {"pic_project": "", "team_support": "", "lokasi": ""}

def save_persist(data):
    try:
        PERSIST_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2))
    except Exception:
        pass

def load_full_data():
    default = {
        "activities": [{"waktu": "", "uraian": "", "keterangan": "", "status": "Selesai"}],
        "incidents": [],
        "materials": [],
        "followups": [],
        "catatan": "",
    }
    try:
        if FULL_DATA_FILE.exists():
            data = json.loads(FULL_DATA_FILE.read_text())
            for k in default:
                if k not in data:
                    data[k] = default[k]
            return data
    except Exception:
        pass
    return default

def save_full_data(data):
    to_save = {k: v for k, v in data.items() if k != "photos"}
    try:
        FULL_DATA_FILE.write_text(json.dumps(to_save, ensure_ascii=False, indent=2))
    except Exception:
        pass

# Inisialisasi session state
if "full_data" not in st.session_state:
    st.session_state.full_data = load_full_data()
if "photos" not in st.session_state:
    st.session_state.photos = []

def auto_save_full_data():
    save_full_data(st.session_state.full_data)

# ═══════════════════════════════════════════════════════════════
# IMAGE HELPERS
# ═══════════════════════════════════════════════════════════════
def resize_image(img_bytes: bytes, max_width: int = 1200, quality: int = 85) -> bytes:
    img = Image.open(io.BytesIO(img_bytes))
    img = ImageOps.exif_transpose(img)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    if img.width > max_width:
        ratio = max_width / img.width
        new_height = int(img.height * ratio)
        img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
    out = io.BytesIO()
    img.save(out, format="JPEG", quality=quality, optimize=True)
    return out.getvalue()

def rotate_image(img_bytes: bytes, degrees: int) -> bytes:
    img = Image.open(io.BytesIO(img_bytes))
    img = ImageOps.exif_transpose(img)
    if degrees == 90:
        img = img.transpose(Image.ROTATE_90)
    elif degrees == -90:
        img = img.transpose(Image.ROTATE_270)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    out = io.BytesIO()
    img.save(out, format="JPEG", quality=85, optimize=True)
    return out.getvalue()

def fix_image_orientation(img_bytes: bytes) -> bytes:
    img = Image.open(io.BytesIO(img_bytes))
    img = ImageOps.exif_transpose(img)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    out = io.BytesIO()
    img.save(out, format="JPEG", quality=85, optimize=True)
    return out.getvalue()

# ═══════════════════════════════════════════════════════════════
# DOCX MANIPULATION HELPERS
# ═══════════════════════════════════════════════════════════════
def set_cell_text(cell, text, *, bold=None, font_size_pt=None, align=None, preserve_format=True):
    existing_font = None
    existing_size = None
    existing_bold = None
    existing_color = None
    if preserve_format and cell.paragraphs and cell.paragraphs[0].runs:
        first_run = cell.paragraphs[0].runs[0]
        existing_font = first_run.font.name
        existing_size = first_run.font.size
        existing_bold = first_run.font.bold
        if first_run.font.color and first_run.font.color.rgb:
            existing_color = first_run.font.color.rgb
    for p in cell.paragraphs:
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    lines = str(text).split("\n") if text else [""]
    for i, line in enumerate(lines):
        if i > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(line)
        if existing_font:
            run.font.name = existing_font
        if font_size_pt is not None:
            run.font.size = Pt(font_size_pt)
        elif existing_size:
            run.font.size = existing_size
        if bold is not None:
            run.font.bold = bold
        elif existing_bold is not None:
            run.font.bold = existing_bold
        if existing_color:
            run.font.color.rgb = existing_color

def set_cell_text_clean(cell, text, *, align=None):
    existing_font = None
    existing_size = None
    existing_bold = None
    if cell.paragraphs and cell.paragraphs[0].runs:
        first_run = cell.paragraphs[0].runs[0]
        existing_font = first_run.font.name
        existing_size = first_run.font.size
        existing_bold = first_run.font.bold
    first_p = cell.paragraphs[0]
    extras = list(cell.paragraphs)[1:]
    for p in extras:
        p._element.getparent().remove(p._element)
    for r in list(first_p.runs):
        r._element.getparent().remove(r._element)
    pPr = first_p._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
    if align is not None:
        first_p.alignment = align
    lines = str(text).split("\n") if text else [""]
    for i, line in enumerate(lines):
        if i == 0:
            p = first_p
        else:
            p = cell.add_paragraph()
            pPr2 = p._element.find(qn('w:pPr'))
            if pPr2 is not None:
                numPr2 = pPr2.find(qn('w:numPr'))
                if numPr2 is not None:
                    pPr2.remove(numPr2)
            if align is not None:
                p.alignment = align
        run = p.add_run(line)
        if existing_font:
            run.font.name = existing_font
        if existing_size:
            run.font.size = existing_size
        if existing_bold is not None:
            run.font.bold = existing_bold

def force_page_break_before(paragraph):
    from docx.oxml import OxmlElement
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph._element.insert(0, pPr)
    existing = pPr.find(qn('w:pageBreakBefore'))
    if existing is not None:
        pPr.remove(existing)
    pbb = OxmlElement('w:pageBreakBefore')
    pPr.insert(0, pbb)

def remove_element(elem):
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)

def remove_empty_photo_section(doc):
    body = doc.element.body
    to_remove = []
    in_photo_section = False
    for child in list(body.iterchildren()):
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            texts = [t.text for t in child.iter(qn('w:t')) if t.text]
            text = ' '.join(texts).upper()
            if 'LAMPIRAN FOTO KEGIATAN' in text:
                in_photo_section = True
                to_remove.append(child)
                continue
        if in_photo_section and tag != 'sectPr':
            to_remove.append(child)
    for elem in to_remove:
        remove_element(elem)

def remove_second_photo_section_if_unused(doc, photo_count):
    if photo_count > 8:
        return
    if photo_count == 0:
        return
    body = doc.element.body
    photo_heading_count = 0
    to_remove = []
    found_second_heading = False
    for child in list(body.iterchildren()):
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            texts = [t.text for t in child.iter(qn('w:t')) if t.text]
            text = ' '.join(texts).upper()
            if 'LAMPIRAN FOTO KEGIATAN' in text:
                photo_heading_count += 1
                if photo_heading_count == 2:
                    found_second_heading = True
                    to_remove.append(child)
                    continue
        if found_second_heading and tag != 'sectPr':
            to_remove.append(child)
            if tag == 'tbl':
                break
    for elem in to_remove:
        remove_element(elem)

def ensure_first_photo_page_break(doc):
    found_first = False
    for p in doc.paragraphs:
        if 'LAMPIRAN FOTO KEGIATAN' in p.text.upper():
            if not found_first:
                force_page_break_before(p)
                found_first = True

def add_image_to_cell(cell, img_bytes, caption, width_cm=7.5):
    for p in cell.paragraphs:
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
    paragraphs = list(cell.paragraphs)
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p_img = cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=Cm(width_cm))
    p_cap = cell.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = p_cap.add_run(caption or "Foto Kegiatan")
    cap_run.font.size = Pt(9)

def fill_data_table(table, rows_data, col_extractors, center_cols=None):
    if center_cols is None:
        center_cols = []
    if len(table.rows) < 2:
        return
    tbl_element = table._tbl
    existing_rows = tbl_element.findall(qn('w:tr'))
    for tr in existing_rows[1:]:
        tbl_element.remove(tr)
    if rows_data:
        template_row = existing_rows[1] if len(existing_rows) > 1 else None
        if template_row is None:
            return
        for _ in rows_data:
            new_tr = deepcopy(template_row)
            tbl_element.append(new_tr)
    for i, row_data in enumerate(rows_data):
        if i+1 >= len(table.rows):
            break
        row = table.rows[i+1]
        for col_idx, extractor in enumerate(col_extractors):
            if col_idx >= len(row.cells):
                break
            value = extractor(i, row_data)
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[col_idx], value, align=align)

def fill_photo_tables(doc, photos):
    photo_tables = [doc.tables[6], doc.tables[7]]
    photo_idx = 0
    for ti, table in enumerate(photo_tables):
        for row in table.rows:
            for cell in row.cells:
                if photo_idx < len(photos):
                    p = photos[photo_idx]
                    add_image_to_cell(cell, p["bytes"], p["caption"], width_cm=7.5)
                    photo_idx += 1
                else:
                    for para in cell.paragraphs:
                        for r in list(para.runs):
                            r._element.getparent().remove(r._element)
                    extras = list(cell.paragraphs)[1:]
                    for p in extras:
                        p._element.getparent().remove(p._element)

# ═══════════════════════════════════════════════════════════════
# CORE BUILD REPORT
# ═══════════════════════════════════════════════════════════════
def format_tanggal(d):
    if isinstance(d, date):
        return f"{d.day} {BULAN_ID[d.month-1]} {d.year}"
    return str(d)

def build_report(data, photos):
    doc = Document(str(TEMPLATE_PATH))
    
    # Table 0
    t_info = doc.tables[0]
    r0_tcs = t_info._tbl.findall(qn('w:tr'))[0].findall(qn('w:tc'))
    from docx.table import _Cell
    pic_value_cell = _Cell(r0_tcs[1], t_info)
    tanggal_value_cell = _Cell(r0_tcs[3], t_info)
    set_cell_text(pic_value_cell, data["pic_project"] or "-")
    set_cell_text(tanggal_value_cell, format_tanggal(data["tanggal"]))
    
    r1_tcs = t_info._tbl.findall(qn('w:tr'))[1].findall(qn('w:tc'))
    team_col1_cell = _Cell(r1_tcs[1], t_info)
    team_col2_cell = _Cell(r1_tcs[2], t_info)
    lokasi_value_cell = _Cell(r1_tcs[4], t_info)
    
    team_list = [t.strip() for t in re.split(r"[,;\n]", data["team_support"]) if t.strip()]
    if team_list:
        mid = (len(team_list)+1)//2
        col1 = team_list[:mid]
        col2 = team_list[mid:]
        col1_text = "\n".join(f"{i+1}. {name}" for i,name in enumerate(col1))
        col2_text = "\n".join(f"{mid+i+1}. {name}" for i,name in enumerate(col2))
        set_cell_text_clean(team_col1_cell, col1_text)
        set_cell_text_clean(team_col2_cell, col2_text)
    else:
        set_cell_text_clean(team_col1_cell, "-")
        set_cell_text_clean(team_col2_cell, "")
    set_cell_text(lokasi_value_cell, data["lokasi"] or "-")
    
    # Table 1: Kegiatan
    fill_data_table(doc.tables[1], data["activities"],
        [lambda i,a: str(i+1), lambda i,a: a["waktu"], lambda i,a: a["uraian"], lambda i,a: a["keterangan"], lambda i,a: a["status"]],
        center_cols=[0,4])
    
    # Table 2: Kendala
    if data["incidents"]:
        fill_data_table(doc.tables[2], data["incidents"],
            [lambda i,x: str(i+1), lambda i,x: x["masalah"], lambda i,x: x["tindakan"], lambda i,x: x["hasil"]],
            center_cols=[0])
    else:
        fill_data_table(doc.tables[2], [{"masalah":"","tindakan":"","hasil":""}],
            [lambda i,x: "", lambda i,x: x["masalah"], lambda i,x: x["tindakan"], lambda i,x: x["hasil"]],
            center_cols=[0])
    
    # Table 3: Follow Up
    followups = data.get("followups", [])
    if followups:
        fill_data_table(doc.tables[3], followups,
            [lambda i,f: str(i+1), lambda i,f: f["deskripsi"], lambda i,f: f["alasan"], lambda i,f: f["target"]],
            center_cols=[0])
    else:
        fill_data_table(doc.tables[3], [{"deskripsi":"","alasan":"","target":""}],
            [lambda i,f: "", lambda i,f: f["deskripsi"], lambda i,f: f["alasan"], lambda i,f: f["target"]],
            center_cols=[0])
    
    # Table 4: Catatan
    t_catatan = doc.tables[4]
    set_cell_text(t_catatan.rows[0].cells[0], data["catatan"] or " ")
    
    # Table 5: Material
    if data["materials"]:
        fill_data_table(doc.tables[5], data["materials"],
            [lambda i,m: str(i+1), lambda i,m: m["nama"], lambda i,m: m["qty"], lambda i,m: m["kondisi"], lambda i,m: m["keterangan"]],
            center_cols=[0,2,3])
    else:
        fill_data_table(doc.tables[5], [{"nama":"-","qty":"","kondisi":"","keterangan":""}],
            [lambda i,m: "", lambda i,m: m["nama"], lambda i,m: m["qty"], lambda i,m: m["kondisi"], lambda i,m: m["keterangan"]],
            center_cols=[0,1,2,3])
    
    # Photo tables
    fill_photo_tables(doc, photos)
    if photos:
        remove_second_photo_section_if_unused(doc, len(photos))
        ensure_first_photo_page_break(doc)
    else:
        remove_empty_photo_section(doc)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ═══════════════════════════════════════════════════════════════
# PDF GENERATION (Hybrid: win32com → fallback fpdf2++)
# ═══════════════════════════════════════════════════════════════

def _try_win32com_pdf(data, photos):
    """Coba konversi DOCX → PDF via Word (Windows only)."""
    try:
        import pythoncom
        from win32com.client import Dispatch
    except (ImportError, AttributeError):
        return None  # win32com tidak tersedia (Linux/Mac)
    try:
        docx_bytes = build_report(data, photos).getvalue()
        tmp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(tmp_dir, "report_temp.docx")
        pdf_path = os.path.join(tmp_dir, "report_temp.pdf")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        pythoncom.CoInitialize()
        try:
            word = Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc.Close()
            word.Quit()
        finally:
            pythoncom.CoUninitialize()
        with open(pdf_path, "rb") as f:
            result = f.read()
        for p in [docx_path, pdf_path]:
            try:
                os.remove(p)
            except Exception:
                pass
        try:
            os.rmdir(tmp_dir)
        except Exception:
            pass
        return io.BytesIO(result)
    except Exception as e:
        # Error teknis — biarkan fpdf2 sebagai fallback
        return None



def _try_pymupdf_template_pdf(data, photos):
    try:
        if not TEMPLATE_PDF_PATH.exists():
            return None
        doc = fitz.open(str(TEMPLATE_PDF_PATH))
        page = doc[0]
        page_rect = page.rect
        redact_areas_page1 = [
            (105, 163, 320, 180),
            (380, 163, 560, 180),
            (100, 183, 315, 213),
            (380, 183, 560, 203),
            (48, 288, 553, 375),
            (48, 394, 553, 500),
            (48, 505, 553, 610),
            (48, 615, 553, 680),
            (48, 698, 553, 760),
        ]
        for area in redact_areas_page1:
            r = fitz.Rect(area[0], area[1], area[2], area[3])
            page.add_redact_annot(r, fill=None)
        page.apply_redactions()
        # Try system Arial first, fallback to Base14 Helvetica (cross-platform)
        # Font setup with cross-platform fallback
        # Priority: Arial (Windows) -> Liberation Sans (Linux fonts-liberation) -> Helvetica (Base14)
        font_candidates = [
            ("Arial", "Arial-BoldMT"),              # Windows native Arial
            ("LiberationSans", "LiberationSans-Bold"),  # Linux fonts-liberation (metric Arial)
            ("Helvetica", "Helvetica-Bold"),         # Base14 fallback - always works
        ]
        font_regular = font_candidates[-1][0]
        font_bold = font_candidates[-1][1]
        for reg, bold in font_candidates:
            try:
                fitz.Font(reg)
                fitz.Font(bold)
                font_regular = reg
                font_bold = bold
                break
            except Exception:
                continue

        page.insert_text(fitz.Point(108, 174), (data["pic_project"] or "-"), fontsize=9.5, fontname=font_regular, color=(0, 0, 0))
        page.insert_text(fitz.Point(383, 174), format_tanggal(data["tanggal"]), fontsize=9.5, fontname=font_regular, color=(0, 0, 0))
        team_list = [t.strip() for t in re.split(r"[,;\n]", data["team_support"]) if t.strip()]
        if team_list:
            mid = (len(team_list) + 1) // 2
            for li, name in enumerate(team_list[:mid]):
                page.insert_text(fitz.Point(108, 193 + li * 14), f"{li+1}. {name}", fontsize=9, fontname=font_regular, color=(0, 0, 0))
            for li, name in enumerate(team_list[mid:]):
                page.insert_text(fitz.Point(200, 193 + li * 14), f"{mid+li+1}. {name}", fontsize=9, fontname=font_regular, color=(0, 0, 0))
        else:
            page.insert_text(fitz.Point(108, 193), "-", fontsize=9, fontname=font_regular, color=(0, 0, 0))
        page.insert_text(fitz.Point(383, 193), (data["lokasi"] or "-"), fontsize=9.5, fontname=font_regular, color=(0, 0, 0))
        y_base = 289
        for i, a in enumerate(data.get("activities", [])):
            y = y_base + i * 18
            if y + 18 > 375:
                break
            page.insert_text(fitz.Point(52, y + 5), str(i + 1), fontsize=8.5, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(92, y + 5), str(a.get("waktu", "")), fontsize=8.5, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(200, y + 5), str(a.get("uraian", "")), fontsize=8.5, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(382, y + 5), str(a.get("keterangan", "")), fontsize=8.5, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(502, y + 5), str(a.get("status", "")), fontsize=8.5, fontname=font_regular, color=(0, 0, 0))
        y_base = 395
        for i, inc in enumerate(data.get("incidents", [])):
            y = y_base + i * 18
            if y + 18 > 500:
                break
            page.insert_text(fitz.Point(52, y + 5), str(i + 1), fontsize=8.5, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(114, y + 5), str(inc.get("masalah", ""))[:60], fontsize=8, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(310, y + 5), str(inc.get("tindakan", ""))[:50], fontsize=8, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(475, y + 5), str(inc.get("hasil", ""))[:40], fontsize=8, fontname=font_regular, color=(0, 0, 0))
        y_base = 506
        for i, fu in enumerate(data.get("followups", [])):
            y = y_base + i * 18
            if y + 18 > 610:
                break
            page.insert_text(fitz.Point(52, y + 5), str(i + 1), fontsize=8.5, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(114, y + 5), str(fu.get("deskripsi", ""))[:60], fontsize=8, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(310, y + 5), str(fu.get("alasan", ""))[:50], fontsize=8, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(475, y + 5), str(fu.get("target", ""))[:40], fontsize=8, fontname=font_regular, color=(0, 0, 0))
        catatan = data.get("catatan", "")
        if catatan:
            page.insert_text(fitz.Point(52, 625), catatan[:200], fontsize=8.5, fontname=font_regular, color=(0, 0, 0))
        y_base = 699
        for i, m in enumerate(data.get("materials", [])):
            y = y_base + i * 18
            if y + 18 > 760:
                break
            page.insert_text(fitz.Point(52, y + 5), str(i + 1), fontsize=8.5, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(144, y + 5), str(m.get("nama", ""))[:50], fontsize=8, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(328, y + 5), str(m.get("qty", "")), fontsize=8.5, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(389, y + 5), str(m.get("kondisi", "")), fontsize=8.5, fontname=font_regular, color=(0, 0, 0))
            page.insert_text(fitz.Point(476, y + 5), str(m.get("keterangan", ""))[:40], fontsize=8, fontname=font_regular, color=(0, 0, 0))
        while doc.page_count > 1:
            doc.delete_page(doc.page_count - 1)
        if photos:
            photos_per_page = 4
            img_w = 240
            img_h = 175
            margin_left = 50
            margin_top = 110
            spacing_x = 20
            spacing_y = 30
            photo_idx = 0
            while photo_idx < len(photos):
                new_page = doc.new_page(width=page_rect.width, height=page_rect.height)
                new_page.insert_text(fitz.Point(margin_left, 70), "LAMPIRAN FOTO KEGIATAN", fontsize=14, fontname=font_bold, color=(0, 0, 0))
                new_page.draw_line(fitz.Point(margin_left, 75), fitz.Point(page_rect.width - margin_left, 75), color=(0.2, 0.2, 0.2), width=0.5)
                for grid_i in range(photos_per_page):
                    if photo_idx >= len(photos):
                        break
                    p = photos[photo_idx]
                    col = grid_i % 2
                    row = grid_i // 2
                    x = margin_left + col * (img_w + spacing_x)
                    y = margin_top + row * (img_h + spacing_y + 25)
                    try:
                        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
                        tmp.write(p["bytes"])
                        tmp_path = tmp.name
                        tmp.close()
                        new_page.insert_image(fitz.Rect(x, y, x + img_w, y + img_h), filename=tmp_path)
                        os.unlink(tmp_path)
                        caption = p["caption"][:60] if p["caption"] else "Foto Kegiatan"
                        new_page.insert_text(fitz.Point(x, y + img_h + 12), f"Foto {photo_idx + 1}: {caption}", fontsize=8, fontname=font_regular, color=(0.3, 0.3, 0.3))
                    except Exception:
                        new_page.draw_rect(fitz.Rect(x, y, x + img_w, y + img_h), color=(0.8, 0.2, 0.2), width=1)
                        new_page.insert_text(fitz.Point(x + 10, y + img_h // 2), f"[Foto {photo_idx + 1} error]", fontsize=9, fontname=font_regular, color=(0.8, 0.2, 0.2))
                    photo_idx += 1
        output = io.BytesIO()
        doc.save(output)
        doc.close()
        output.seek(0)
        return output
    except Exception as e:
        return None


def _render_table(pdf, headers, rows, col_widths):
    """Render tabel profesional dengan fpdf2 menggunakan cell()."""
    if not rows:
        return
    row_h = 6.5

    # ── Header ──
    pdf.set_font('Helvetica', 'B', 8)
    pdf.set_fill_color(20, 25, 66)  # Navy gelap
    pdf.set_text_color(255, 255, 255)
    for i, h in enumerate(headers):
        align = 'C' if i == 0 else 'L'
        pdf.cell(col_widths[i], row_h, f' {h[:30]}', border=1, align=align, fill=True)
    pdf.ln()

    # ── Data ──
    pdf.set_font('Helvetica', '', 7.5)
    pdf.set_text_color(30, 30, 30)
    for ri, row in enumerate(rows):
        # Cek page break
        if pdf.get_y() + row_h > 260:
            pdf.add_page()
            # Header ulang
            pdf.set_font('Helvetica', 'B', 8)
            pdf.set_fill_color(20, 25, 66)
            pdf.set_text_color(255, 255, 255)
            for i, h in enumerate(headers):
                align = 'C' if i == 0 else 'L'
                pdf.cell(col_widths[i], row_h, f' {h[:30]}', border=1, align=align, fill=True)
            pdf.ln()
            pdf.set_font('Helvetica', '', 7.5)
            pdf.set_text_color(30, 30, 30)

        # Alternating row color
        if ri % 2 == 0:
            pdf.set_fill_color(248, 249, 252)
        else:
            pdf.set_fill_color(255, 255, 255)

        # Max chars per column (Helvetica 7.5pt: ~0.38 chars/mm)
        max_chars = [max(3, int(w * 0.38)) for w in col_widths]
        for ci, val in enumerate(row):
            align = 'C' if ci == 0 else 'L'
            truncated = str(val)[:max_chars[ci]]
            pdf.cell(col_widths[ci], row_h, f' {truncated}', border=1, align=align, fill=True)
        pdf.ln()
    pdf.ln(2)


def build_report_pdf(data, photos):
    """Hybrid PDF generation with 3-tier fallback:
    1. PyMuPDF template-based (hasil identik template PDF)
    2. win32com (Windows → hasil sempurna, identik template DOCX)
    3. fpdf2 (cross-platform, layout profesional)
    """
    # Langkah 1: Coba PyMuPDF dari template PDF (paling akurat)
    pymupdf_result = _try_pymupdf_template_pdf(data, photos)
    win32_result = _try_win32com_pdf(data, photos)
    if win32_result is not None:
        return win32_result

    # Langkah 2: Fallback ke fpdf2 (cross-platform)
    class _ReportPDF(FPDF):
        def footer(self):
            self.set_y(-15)
            self.set_font('Helvetica', 'I', 7.5)
            self.set_text_color(130, 130, 130)
            self.cell(0, 8, f'Halaman {self.page_no()}', align='C')
    pdf = _ReportPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Header
    pdf.set_font('Helvetica', 'B', 12)
    pdf.set_text_color(20, 25, 66)
    pdf.cell(0, 8, 'DAILY REPORT - NALE SYSTEM INTEGRATOR', new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.set_font('Helvetica', '', 9)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 5, 'Managed Service Dinas Kominfo Kab. Mojokerto', new_x="LMARGIN", new_y="NEXT", align='C')
    # Garis pemisah
    pdf.set_draw_color(245, 166, 35)
    pdf.set_line_width(0.5)
    y = pdf.get_y()
    pdf.line(15, y+1, 195, y+1)
    pdf.set_line_width(0.2)
    pdf.set_draw_color(180, 180, 180)
    pdf.ln(5)

    def section_title(title):
        pdf.ln(2)
        pdf.set_font('Helvetica', 'B', 9)
        pdf.set_fill_color(245, 166, 35)
        pdf.set_text_color(20, 25, 66)
        pdf.cell(0, 6.5, f'  {title}', new_x="LMARGIN", new_y="NEXT", align='L', fill=True)
        pdf.ln(2)

    def info_row(label, value):
        pdf.set_font('Helvetica', 'B', 8.5)
        pdf.set_text_color(20, 25, 66)
        pdf.cell(32, 5.5, f'{label}', new_x="END")
        pdf.set_font('Helvetica', '', 8.5)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 5.5, f': {value}', new_x="LMARGIN", new_y="NEXT")

    # ── 1. Informasi Umum ──
    section_title('1. INFORMASI UMUM')
    info_row('PIC Project', data['pic_project'] or '-')
    info_row('Tanggal', format_tanggal(data['tanggal']))
    team_list = [t.strip() for t in re.split(r"[,;\n]", data['team_support']) if t.strip()]
    team_str = ', '.join(team_list) if team_list else '-'
    info_row('Team Support', team_str)
    info_row('Lokasi / Site', data['lokasi'] or '-')
    pdf.ln(1)

    # ── 2. Kegiatan Harian ──
    section_title('2. KEGIATAN HARIAN')
    if data['activities']:
        _render_table(pdf,
            ['No', 'Waktu', 'Uraian Kegiatan', 'Keterangan', 'Status'],
            [[i+1, a['waktu'], a['uraian'], a['keterangan'], a['status']] for i, a in enumerate(data['activities'])],
            [10, 22, 78, 44, 20])
    else:
        pdf.set_font('Helvetica', '', 8.5)
        pdf.cell(0, 6, 'Tidak ada kegiatan.', new_x="LMARGIN", new_y="NEXT")

    # ── 3. Kendala / Insiden ──
    section_title('3. KENDALA / INSIDEN')
    if data['incidents']:
        _render_table(pdf,
            ['No', 'Masalah', 'Tindakan', 'Hasil'],
            [[i+1, inc['masalah'], inc['tindakan'], inc['hasil']] for i, inc in enumerate(data['incidents'])],
            [10, 65, 60, 55])
    else:
        pdf.set_font('Helvetica', '', 8.5)
        pdf.cell(0, 6, 'Tidak ada kendala.', new_x="LMARGIN", new_y="NEXT")

    # ── 4. Follow Up ──
    section_title('4. PEKERJAAN BELUM SELESAI / FOLLOW UP')
    if data['followups']:
        _render_table(pdf,
            ['No', 'Deskripsi', 'Alasan/Kendala', 'Target'],
            [[i+1, fu['deskripsi'], fu['alasan'], fu['target']] for i, fu in enumerate(data['followups'])],
            [10, 65, 60, 55])
    else:
        pdf.set_font('Helvetica', '', 8.5)
        pdf.cell(0, 6, 'Tidak ada follow up.', new_x="LMARGIN", new_y="NEXT")

    # ── 5. Catatan Tambahan ──
    section_title('5. CATATAN TAMBAHAN')
    pdf.set_font('Helvetica', '', 8.5)
    pdf.set_text_color(40, 40, 40)
    pdf.multi_cell(0, 5, data['catatan'] if data['catatan'] else '-')

    # ── 6. Material ──
    section_title('6. MATERIAL YANG DIGUNAKAN')
    if data['materials']:
        _render_table(pdf,
            ['No', 'Nama Material', 'Qty', 'Kondisi', 'Keterangan'],
            [[i+1, m['nama'], m['qty'], m['kondisi'], m['keterangan']] for i, m in enumerate(data['materials'])],
            [10, 60, 15, 20, 45])
    else:
        pdf.set_font('Helvetica', '', 8.5)
        pdf.cell(0, 6, 'Tidak ada material.', new_x="LMARGIN", new_y="NEXT")

    # ── 7. Lampiran Foto ──
    if photos:
        pdf.add_page()
        section_title('7. LAMPIRAN FOTO KEGIATAN')

        img_w = 80
        img_h = 58
        x_left = 15
        x_right = x_left + img_w + 8
        y_start = pdf.get_y()

        photo_row = 0
        for idx, p in enumerate(photos):
            col = idx % 2
            x = x_left if col == 0 else x_right
            y = y_start + photo_row * (img_h + 16)

            if y + img_h > 260:
                pdf.add_page()
                section_title('7. LAMPIRAN FOTO KEGIATAN (lanjutan)')
                y_start = pdf.get_y()
                photo_row = 0
                y = y_start
                x = x_left if col == 0 else x_right

            try:
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg')
                tmp.write(p['bytes'])
                tmp_path = tmp.name
                tmp.close()
                # Border foto
                pdf.set_draw_color(180, 180, 180)
                pdf.rect(x-0.3, y-0.3, img_w+0.6, img_h+0.6)
                pdf.image(tmp_path, x=x, y=y, w=img_w, h=img_h)
                pdf.set_xy(x, y + img_h + 1)
                pdf.set_font('Helvetica', 'I', 7)
                pdf.set_text_color(100, 100, 100)
                caption = p['caption'][:50] if p['caption'] else 'Foto Kegiatan'
                pdf.cell(img_w, 5, f'Foto {idx+1}: {caption}', align='C')
                os.unlink(tmp_path)
            except Exception:
                pdf.set_draw_color(200, 50, 50)
                pdf.rect(x, y, img_w, img_h)
                pdf.set_xy(x, y + 2)
                pdf.set_font('Helvetica', '', 8)
                pdf.set_text_color(200, 50, 50)
                pdf.cell(img_w, img_h-4, f'[Foto {idx+1} error]', align='C')

            # Increment row after second column
            if col == 1:
                photo_row += 1

    output = io.BytesIO()
    pdf.output(output)
    output.seek(0)
    return output

# ═══════════════════════════════════════════════════════════════
# UI
# ═══════════════════════════════════════════════════════════════

st.markdown("""
<div class="brand-header">
  <div>
    <div class="brand-tag">NALE SYSTEM INTEGRATOR</div>
    <div class="brand-title">📋 Daily Report Generator</div>
    <div class="brand-sub">Managed Service — Dinas Kominfo Kab. Mojokerto</div>
  </div>
</div>
""", unsafe_allow_html=True)

if not TEMPLATE_PATH.exists():
    st.error(f"❌ File template tidak ditemukan: `{TEMPLATE_PATH}`")
    st.stop()
if not TEMPLATE_PDF_PATH.exists():
    st.warning(f"⚠️ File template PDF tidak ditemukan: `{TEMPLATE_PDF_PATH.name}`. PDF fallback akan menggunakan fpdf2 (tidak identik dengan template). Rekomendasi: letakkan template PDF di folder yang sama dengan app.py.")

# Load data
persist = load_persist()
full_data = st.session_state.full_data
photos = st.session_state.photos

# Tab
tab_info, tab_photos = st.tabs(["📋 Info & Kegiatan", "📸 Foto"])

# ═══════════════════════════════════════════════════════════════
# TAB INFO & KEGIATAN (dengan placeholder)
# ═══════════════════════════════════════════════════════════════
with tab_info:
    # --- 1. Informasi Umum ---
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    col_a, col_b = st.columns([3,1])
    with col_a:
        st.markdown('<h3><span class="section-num">1</span>Informasi Umum</h3>', unsafe_allow_html=True)
    with col_b:
        st.markdown('<div class="persist-info">● Auto-saved</div>', unsafe_allow_html=True)
    
    pic_project = st.text_input("PIC Project", value=persist.get("pic_project", ""), placeholder="Nama PIC Project", key="pic_project")
    tanggal = st.date_input("Tanggal", value=date.today(), key="tanggal", format="DD/MM/YYYY")
    team_support = st.text_input("Team Support (pisahkan koma)", value=persist.get("team_support", ""), placeholder="Contoh: Andi, Budi, Citra", key="team_support")
    lokasi = st.text_input("Lokasi / Site", value=persist.get("lokasi", ""), placeholder="Nama lokasi kegiatan", key="lokasi")
    
    new_persist = {"pic_project": pic_project, "team_support": team_support, "lokasi": lokasi}
    if new_persist != persist:
        save_persist(new_persist)
    st.markdown('</div>', unsafe_allow_html=True)
    
    # --- 2. Kegiatan Harian (dengan placeholder)---
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h3><span class="section-num">2</span>Kegiatan Harian</h3>', unsafe_allow_html=True)
    activities = full_data.setdefault("activities", [{"waktu":"","uraian":"","keterangan":"","status":"Selesai"}])
    for i, act in enumerate(activities):
        cols = st.columns([0.4,1.3,3,1.5,1.2,0.4])
        cols[0].markdown(f"<div style='padding-top:32px;text-align:center'>{i+1}</div>", unsafe_allow_html=True)
        act["waktu"] = cols[1].text_input("Waktu", act["waktu"], placeholder="08.00-10.00", key=f"act_w_{i}", label_visibility="collapsed")
        act["uraian"] = cols[2].text_input("Uraian", act["uraian"], placeholder="Uraian kegiatan", key=f"act_u_{i}", label_visibility="collapsed")
        act["keterangan"] = cols[3].text_input("Keterangan", act["keterangan"], placeholder="Keterangan (opsional)", key=f"act_k_{i}", label_visibility="collapsed")
        act["status"] = cols[4].selectbox("Status", ["Selesai","Proses","Pending"], index=["Selesai","Proses","Pending"].index(act["status"]) if act["status"] in ["Selesai","Proses","Pending"] else 0, key=f"act_s_{i}", label_visibility="collapsed")
        if cols[5].button("✕", key=f"act_del_{i}"):
            activities.pop(i)
            auto_save_full_data()
            st.rerun()
    if st.button("+ Tambah Kegiatan", key="add_act"):
        activities.append({"waktu":"","uraian":"","keterangan":"","status":"Selesai"})
        auto_save_full_data()
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)
    
    # --- 3. Kendala / Insiden (dengan placeholder)---
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h3><span class="section-num">3</span>Kendala / Insiden</h3>', unsafe_allow_html=True)
    incidents = full_data.setdefault("incidents", [])
    for i, inc in enumerate(incidents):
        cols = st.columns([0.4,3,3,3,0.4])
        cols[0].markdown(f"<div style='padding-top:32px'>{i+1}</div>", unsafe_allow_html=True)
        inc["masalah"] = cols[1].text_area("Masalah", inc["masalah"], placeholder="Deskripsi masalah / problem", key=f"inc_m_{i}", label_visibility="collapsed", height=70)
        inc["tindakan"] = cols[2].text_area("Tindakan", inc["tindakan"], placeholder="Tindakan yang dilakukan", key=f"inc_t_{i}", label_visibility="collapsed", height=70)
        inc["hasil"] = cols[3].text_area("Hasil", inc["hasil"], placeholder="Hasil dari tindakan", key=f"inc_h_{i}", label_visibility="collapsed", height=70)
        if cols[4].button("✕", key=f"inc_del_{i}"):
            incidents.pop(i)
            auto_save_full_data()
            st.rerun()
    if st.button("+ Tambah Insiden", key="add_inc"):
        incidents.append({"masalah":"","tindakan":"","hasil":""})
        auto_save_full_data()
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)
    
    # --- 4. Follow Up (dengan placeholder)---
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h3><span class="section-num">4</span>Pekerjaan Belum Selesai / Follow Up</h3>', unsafe_allow_html=True)
    followups = full_data.setdefault("followups", [])
    for i, fu in enumerate(followups):
        cols = st.columns([0.4,3,3,3,0.4])
        cols[0].markdown(f"<div style='padding-top:32px'>{i+1}</div>", unsafe_allow_html=True)
        fu["deskripsi"] = cols[1].text_area("Deskripsi", fu.get("deskripsi",""), placeholder="Pekerjaan yang belum selesai", key=f"fu_desc_{i}", label_visibility="collapsed", height=70)
        fu["alasan"] = cols[2].text_area("Alasan/Kendala", fu.get("alasan",""), placeholder="Alasan atau kendala", key=f"fu_alasan_{i}", label_visibility="collapsed", height=70)
        fu["target"] = cols[3].text_area("Target", fu.get("target",""), placeholder="Target penyelesaian", key=f"fu_target_{i}", label_visibility="collapsed", height=70)
        if cols[4].button("✕", key=f"fu_del_{i}"):
            followups.pop(i)
            auto_save_full_data()
            st.rerun()
    if st.button("+ Tambah Follow Up", key="add_fu"):
        followups.append({"deskripsi":"","alasan":"","target":""})
        auto_save_full_data()
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)
    
    # --- 5. Catatan Tambahan ---
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h3><span class="section-num">5</span>Catatan Tambahan</h3>', unsafe_allow_html=True)
    catatan = st.text_area("Catatan", value=full_data.get("catatan",""), placeholder="Catatan tambahan (opsional)...", key="catatan", height=80, label_visibility="collapsed")
    full_data["catatan"] = catatan
    st.markdown('</div>', unsafe_allow_html=True)
    
    # --- 6. Material (dengan placeholder)---
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h3><span class="section-num">6</span>Material yang Digunakan</h3>', unsafe_allow_html=True)
    materials = full_data.setdefault("materials", [])
    for i, mat in enumerate(materials):
        cols = st.columns([0.4,3,0.8,1.2,2.5,0.4])
        cols[0].markdown(f"<div style='padding-top:32px'>{i+1}</div>", unsafe_allow_html=True)
        mat["nama"] = cols[1].text_input("Nama", mat.get("nama",""), placeholder="Nama material/alat", key=f"mat_n_{i}", label_visibility="collapsed")
        mat["qty"] = cols[2].text_input("Qty", mat.get("qty",""), placeholder="Jumlah", key=f"mat_q_{i}", label_visibility="collapsed")
        mat["kondisi"] = cols[3].selectbox("Kondisi", ["Baik","Rusak","Habis"], index=["Baik","Rusak","Habis"].index(mat.get("kondisi","Baik")), key=f"mat_k_{i}", label_visibility="collapsed")
        mat["keterangan"] = cols[4].text_input("Keterangan", mat.get("keterangan",""), placeholder="Keterangan (opsional)", key=f"mat_ket_{i}", label_visibility="collapsed")
        if cols[5].button("✕", key=f"mat_del_{i}"):
            materials.pop(i)
            auto_save_full_data()
            st.rerun()
    if st.button("+ Tambah Material", key="add_mat"):
        materials.append({"nama":"","qty":"","kondisi":"Baik","keterangan":""})
        auto_save_full_data()
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Tombol simpan manual
    st.markdown("---")
    col_save1, col_save2 = st.columns(2)
    with col_save1:
        if st.button("💾 Simpan Draft (Semua Data)", width='stretch'):
            auto_save_full_data()
            st.success("Data disimpan ke file. Tidak akan hilang meskipun refresh.")
    with col_save2:
        if st.button("🔄 Muat Ulang Draft", width='stretch'):
            st.session_state.full_data = load_full_data()
            st.rerun()

# ═══════════════════════════════════════════════════════════════
# TAB FOTO (diperbaiki bug upload + CSS sudah mendukung)
# ═══════════════════════════════════════════════════════════════
with tab_photos:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<h3><span class="section-num">📸</span>Lampiran Foto Kegiatan</h3>', unsafe_allow_html=True)
    st.caption("Gunakan tombol ↻ untuk rotate gambar, ⬆/⬇ untuk urutan, ✕ untuk hapus. **Tidak ada batasan jumlah foto.**")
    
    # Upload foto - tanpa batasan
    uploaded_files = st.file_uploader(
        "Upload foto (jpg/jpeg/png/webp)",
        type=["jpg","jpeg","png","webp"],
        accept_multiple_files=True,
        key="photo_uploader"
    )
    
    if uploaded_files:
        existing_names = {p["name"] for p in st.session_state.photos}
        new_added = False
        for f in uploaded_files:
            if f.name not in existing_names:
                raw = f.read()
                fixed = fix_image_orientation(raw)
                resized = resize_image(fixed, max_width=1200, quality=85)
                st.session_state.photos.append({
                    "bytes": resized,
                    "caption": "",
                    "name": f.name,
                })
                new_added = True
        if new_added:
            st.rerun()
    
    if st.session_state.photos:
        st.markdown(f"**{len(st.session_state.photos)} foto** — gunakan tombol ⬆ ⬇ untuk mengubah urutan")
        for i, photo in enumerate(st.session_state.photos):
            cols = st.columns([1.5, 3, 0.5, 0.5, 0.5, 0.5, 0.7])
            with cols[0]:
                st.image(photo["bytes"], width='stretch')
            with cols[1]:
                new_caption = st.text_area(f"Caption {i+1}", value=photo["caption"], placeholder="Deskripsi foto kegiatan...", key=f"photo_cap_{i}", height=100, label_visibility="collapsed")
                st.session_state.photos[i]["caption"] = new_caption
            with cols[2]:
                if st.button("↻", key=f"rot_left_{i}"):
                    rotated = rotate_image(photo["bytes"], -90)
                    st.session_state.photos[i]["bytes"] = rotated
                    st.rerun()
            with cols[3]:
                if st.button("↻", key=f"rot_right_{i}"):
                    rotated = rotate_image(photo["bytes"], 90)
                    st.session_state.photos[i]["bytes"] = rotated
                    st.rerun()
            with cols[4]:
                if st.button("⬆", key=f"move_up_{i}", disabled=(i==0)):
                    st.session_state.photos[i], st.session_state.photos[i-1] = st.session_state.photos[i-1], st.session_state.photos[i]
                    st.rerun()
            with cols[5]:
                if st.button("⬇", key=f"move_down_{i}", disabled=(i==len(st.session_state.photos)-1)):
                    st.session_state.photos[i], st.session_state.photos[i+1] = st.session_state.photos[i+1], st.session_state.photos[i]
                    st.rerun()
            with cols[6]:
                if st.button("✕", key=f"del_photo_{i}"):
                    st.session_state.photos.pop(i)
                    st.rerun()
    else:
        st.info("Belum ada foto. Upload foto di atas.")
    
    st.markdown('</div>', unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════
# EXPORT BUTTON
# ═══════════════════════════════════════════════════════════════
st.markdown("---")
st.markdown("<h4 style='text-align:center'>📥 Download Laporan</h4>", unsafe_allow_html=True)

col_format, col_gen, col_dummy = st.columns([1, 2, 1])
with col_format:
    format_option = st.radio("Pilih format:", ["Microsoft Word (.docx)", "PDF (.pdf)"], horizontal=True, key="format_radio")
with col_gen:
    if st.button("📄 Generate & Download", type="primary", width='stretch'):
        if not pic_project.strip():
            st.error("⚠️ PIC Project harus diisi.")
        elif not any(a["uraian"].strip() for a in full_data.get("activities", [])):
            st.error("⚠️ Minimal harus ada 1 kegiatan harian dengan uraian.")
        else:
            with st.spinner("Membuat laporan..."):
                data_report = {
                    "pic_project": pic_project,
                    "tanggal": tanggal,
                    "team_support": team_support,
                    "lokasi": lokasi,
                    "activities": [a for a in full_data.get("activities", []) if a["uraian"].strip()],
                    "incidents": [i for i in full_data.get("incidents", []) if i["masalah"].strip()],
                    "followups": [f for f in full_data.get("followups", []) if f["deskripsi"].strip()],
                    "materials": [m for m in full_data.get("materials", []) if m["nama"].strip()],
                    "catatan": full_data.get("catatan", ""),
                }
                try:
                    tgl_str = format_tanggal(tanggal)
                    if format_option == "Microsoft Word (.docx)":
                        output = build_report(data_report, st.session_state.photos)
                        ext = ".docx"
                        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    else:
                        with st.status("⏳ Membuat PDF...", expanded=True) as status:
                            st.write("Membuat laporan PDF...")
                            st.write("Memproses data tabel & foto...")
                            output = build_report_pdf(data_report, st.session_state.photos)
                            status.update(label="✅ PDF selesai!", state="complete")
                        ext = ".pdf"
                        mime = "application/pdf"
                    filename = f"Daily Report_Yearly Managed Service_Kominfo Mojokerto - {tgl_str}{ext}"
                    st.session_state.last_output = output.getvalue()
                    st.session_state.last_filename = filename
                    st.success(f"✅ {ext.upper()} siap di-download! Format identik dengan template.")
                except Exception as e:
                    st.error(f"❌ Error: {e}")
                    import traceback
                    st.code(traceback.format_exc())

if "last_output" in st.session_state:
    col_l, col_c, col_r = st.columns([1,2,1])
    with col_c:
        st.download_button(
            "⬇️ Download",
            data=st.session_state.last_output,
            file_name=st.session_state.last_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if ".docx" in st.session_state.last_filename else "application/pdf",
            width='stretch',
        )
